import subprocess
from ContentGenerators.EssayGenerator.EssayThinker import GenerateBodyParagraphIdea
from ContentGenerators.EssayGenerator.SeperateFunctions import *
from ContentGenerators.GlobalTools.FileManager import put_watermark
from ContentGenerators.GlobalTools.ImageDownloader import download
from ContentGenerators.GlobalTools.Requestor import reply_to
import os
import shutil
from docx.shared import Inches
import docx
from ContentGenerators.GlobalTools.Translator import translate
from docx2pdf import convert


def sort(d):
    return dict(sorted(d.items(), key=lambda item: int(item[0])))

def generate_pdf(doc_path, path):
    subprocess.call(['soffice',
                     # '--headless',
                     '--convert-to',
                     'pdf',
                     '--outdir',
                     path,
                     doc_path])
    return doc_path

class Document:
    def __init__(self,
              id=None,
              path=".", #
              topic="The Evolution of Artificial Intelligence.",
              commands="",
              paragraphs=5,
              wordcount=1000,
              pictures=False,
              language="en",
              search_language="en",
              citation_format="MLA"
              ):
        if id is None: self.id = randint(1, int("9"*10))
        else:
            self.id = id
        self.path = path
        self.topic = topic
        self.commands = commands
        self.paragraph_count = paragraphs
        self.wordcount = wordcount
        self.pictures = pictures
        self.paragraphs = {}
        self.language = language
        self.search_language = search_language
        self.citation_format = citation_format
        self.type = "essay"

        self.filenametitle = f"GeneratedEssay_{self.id}"
        self.filetype = "docx"
        self.uid = None
        self.token = None
        self.hashstring = None
        self.data = None
        self.picture_list = None
        self.citations = {}
        self.history = None

    def Generate(self):
        # Janky, but bear with me here. 😅
        paragraphs = self.paragraph_count
        wordcount = self.wordcount
        lang = self.language
        searchlang = self.search_language
        citationformat = self.citation_format
        hashstring = self.id
        command = self.commands
        topic = translate(self.topic, 'en', lang)
        pictureput = self.pictures
        previousimagelinks = ""
        wordpp = wordcount / paragraphs
        wordpp = int(round(wordpp))  # WORD-PER-PARAGRAPH, THIS IS HOW WE INSURE ACCURACY WITH SELECTED WORDCOUNT.
        wordpp = wordpp
        originaltopic = topic
        citations = ""
        currentparagraphs = 2


        if command.lower() == "nothing" or command == "" or command == " ":
            adins = ""
        else:
            try:
                command = translate(command, 'en', lang)
                adins = "AND PROVIDE INSTRUCTIONS THAT STATISFY THE USER'S COMMAND BUT KEEP EVERYTHING IN ONE SIMPLE 20 WORD SENTENCE: '" + command + "'"
            except:
                adins = ""

        evalidealist = GenerateBodyParagraphIdea(paragraphs, topic, adins)
        idealist = str(evalidealist)

        if pictureput:
            self.picture_list = GeneratePictures(evalidealist, topic, searchlang, lang, previousimagelinks)

        for i in range(1, paragraphs):
            if currentparagraphs == paragraphs:

                concidea = reply_to(prompt.Create("essay", "conclusion_idea", {
                    "topic": topic,
                    "body_paragraphs": str(paragraphs - 2),
                    "idealist": idealist,
                    "adins": adins
                }))

                tries = 0
                while len(concidea) > 220 or idealist.find(
                        concidea) >= 0 or concidea == "Unable to fetch the response, Please try again." and tries < 10:
                    concidea = reply_to(prompt.Create("essay", "conclusion_idea", {
                        "topic": topic,
                        "body_paragraphs": str(paragraphs - 2),
                        "idealist": idealist,
                        "adins": adins
                    }))
                    tries = tries + 1

                # GENERATE A CONCLUSION ACCORDING TO A NEW IDEA AND THE PREVIOUS IDEAS. AKA WRAP-UP
                conclusion = reply_to(prompt.Create("essay", "conclusion", {
                    "wordpp": str(wordpp),
                    "concidea": concidea,
                    "topic": topic
                }))

                conclusion = conclusion.replace("\n", " ").encode('utf-8').decode('unicode_escape')
                self.paragraphs[paragraphs-1] = conclusion
                # GENERATE THE INTRODUCTION ACCORDING THE PREVIOUS IDEAS
                introduction = reply_to(prompt.Create("essay", "introduction", {
                    "wordpp": str(wordpp),
                    "topic": topic,
                    "idealist": idealist,
                    "command": command
                }))

                introduction = introduction.replace("\n", " ").encode('utf-8').decode('unicode_escape')
                self.paragraphs[0] = introduction

                print("Introduction: " + introduction)
                print("Conclusion: " + conclusion)
                if pictureput == True:
                    filename = str(randint(1, 99999999999999999999999999999999999999999999999999))
                    self.picture_list[0] = filename
                    picturetopic = translate(topic, searchlang, lang).replace(".", "").lower()
                    try:
                        download(previousimagelinks, filename,
                                 translate(picturetopic, searchlang, lang),
                                 limit=1,
                                 output_dir='', adult_filter_off=False,
                                 force_replace=False, timeout=60, verbose=True)
                    except:
                        "PICTURE WASNT DOWNLOADED."
            else:
                idea = evalidealist[currentparagraphs]  # Get the idea
                oldidea = evalidealist[currentparagraphs - 1]  # Get the idea of the previous paragraph
                print(f"Generating Paragraph {currentparagraphs}")
                paragraph, citation = GenerateParagraph(searchlang, lang, topic, originaltopic, citationformat, idea, oldidea, wordpp)
                self.paragraphs[currentparagraphs-1] = paragraph
                self.citations[currentparagraphs-1] = citation
                # After all the information has been obtained, a thread for writing the essay paragraph will be started
                # meaning all paragraphs will be generating at the same time.
            currentparagraphs = currentparagraphs + 1
        for i in range(1, len(self.citations) + 1):
            # Citation sorter from first to last.
            citation = self.citations[i]
            citations = citations + "\nBody Paragraph " + str(i) + ":\n" + citation + "\n"
        self.citations = citations
        self.paragraphs = sort(self.paragraphs)


    def Assemble(self, pdf=True, watermark=False):
        self.filetype = "docx"
        if pdf: self.filetype = "pdf"
        topic = self.topic
        paragraphs = self.paragraph_count
        lang = self.language
        pictureput = self.pictures
        pads = self.paragraphs
        picd = self.picture_list
        path = self.path
        justify = True
        citations = self.citations
        filenametitle = self.filenametitle
        filetype = self.filetype

        documentpath = path + "/"
        check = 0
        picnum = 0
        doc = docx.Document()
        doc.add_paragraph(translate(topic, lang, lang), style='Heading 1')
        for e in pads:
            texttoadd = pads[e]
            # ADD THE PARAGRAPH:
            doc.add_paragraph("      " + texttoadd.replace("â", ""), style="Normal")
            # THEN ADD THE PICTURE AFTER IT (AFTER THE USER HAS ASKED FOR IT)
            if e != (paragraphs - 1):
                if pictureput == True:
                    tosearch = picd[picnum]
                    filnamesec = "Image_1"
                    for format in ["jpeg", "jpg", "png"]:
                        try:
                            path = f"{tosearch}/{filnamesec}.{format}"
                            doc.add_picture(path, width=Inches(6.0))
                            break
                        except Exception as e:
                            print(e)
                    try:
                        shutil.rmtree(tosearch)
                    except:
                        "COULDN'T DELETE"
                    picnum = picnum + 1
            # Change the font to Times New Roman, size 14
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if check == 0:
                    run.font.name = 'Times New Roman'
                    run.font.size = docx.shared.Pt(30)
                else:
                    run.font.name = 'Times New Roman'
                    run.font.size = docx.shared.Pt(14)
                check = check + 1
            if justify:
                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.space_before = docx.shared.Pt(10)
        # IF THERE WAS A SEARCH FAILURE, THEN WE HAVE TO MAKE SURE THERE ISNT A SPECIFIED AREA FOR CITED RESOURCES WITH
        # NOTHING UNDER IT.
        if citations == "" or citations is None:
            "NO RESOURCES."
        else:
            citations = "Resource Citations: " + citations
        doc.add_paragraph(citations, style="Normal")
        # Save the document and display it to the user.
        file = documentpath + filenametitle + '.docx'
        doc.save(file)

        if filetype.lower() == "pdf":
            try:
                convert(file)
            except:
                generate_pdf(file, documentpath)
            os.remove(file)
            file = file.replace(".docx", ".pdf")
            if watermark:
                put_watermark(file)
        return file
